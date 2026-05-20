"""Parsers for PDF / DOCX / Markdown / TXT.

Each parser returns a list of ``ParsedSection``s. Sections preserve page
numbers (PDF) and headings (Markdown) so retrieval can show useful sources.
"""

from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path


@dataclass
class ParsedSection:
    text: str
    page_number: int | None = None
    section_title: str | None = None
    meta: dict[str, object] = field(default_factory=dict)


SUPPORTED_EXTS = {".pdf", ".docx", ".md", ".markdown", ".txt"}


def supported_extensions() -> set[str]:
    return SUPPORTED_EXTS


def parse_file(path: Path) -> list[ParsedSection]:
    ext = path.suffix.lower()
    if ext == ".pdf":
        return _parse_pdf(path)
    if ext == ".docx":
        return _parse_docx(path)
    if ext in {".md", ".markdown"}:
        return _parse_markdown(path)
    if ext == ".txt":
        return _parse_txt(path)
    raise ValueError(f"unsupported file type: {ext}")


def _parse_pdf(path: Path) -> list[ParsedSection]:
    # Imported lazily because pymupdf has a heavy native init.
    import pymupdf

    sections: list[ParsedSection] = []
    with pymupdf.open(path) as doc:
        for page_index, page in enumerate(doc):
            text = page.get_text("text").strip()
            if not text:
                continue
            sections.append(ParsedSection(text=text, page_number=page_index + 1))
    return sections


def _parse_docx(path: Path) -> list[ParsedSection]:
    import docx

    document = docx.Document(str(path))
    paragraphs: list[str] = []
    current_heading: str | None = None
    sections: list[ParsedSection] = []
    for para in document.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name if para.style else ""
        if style.startswith("Heading"):
            if paragraphs:
                sections.append(
                    ParsedSection(
                        text="\n\n".join(paragraphs),
                        section_title=current_heading,
                    )
                )
                paragraphs = []
            current_heading = text
        else:
            paragraphs.append(text)
    if paragraphs:
        sections.append(ParsedSection(text="\n\n".join(paragraphs), section_title=current_heading))
    if not sections:
        # Document had no paragraph-level structure — fall back to the whole text.
        whole = "\n".join(p.text for p in document.paragraphs).strip()
        if whole:
            sections.append(ParsedSection(text=whole))
    return sections


def _parse_markdown(path: Path) -> list[ParsedSection]:
    # markdown-it-py token stream → recover headings + plain text.
    from markdown_it import MarkdownIt

    md = MarkdownIt("commonmark")
    raw = path.read_text(encoding="utf-8", errors="replace")
    tokens = md.parse(raw)

    sections: list[ParsedSection] = []
    current_heading: str | None = None
    current_text: list[str] = []
    i = 0
    while i < len(tokens):
        tok = tokens[i]
        if tok.type == "heading_open":
            if current_text:
                sections.append(
                    ParsedSection(
                        text="\n\n".join(s.strip() for s in current_text if s.strip()),
                        section_title=current_heading,
                    )
                )
                current_text = []
            # heading text is in the next inline token
            inline = tokens[i + 1] if i + 1 < len(tokens) else None
            current_heading = inline.content.strip() if inline and inline.content else None
            i += 3  # heading_open, inline, heading_close
            continue
        if tok.type == "inline":
            current_text.append(tok.content)
        i += 1
    if current_text:
        sections.append(
            ParsedSection(
                text="\n\n".join(s.strip() for s in current_text if s.strip()),
                section_title=current_heading,
            )
        )
    if not sections:
        # No headings — emit the whole file as one section.
        text = raw.strip()
        if text:
            sections.append(ParsedSection(text=text))
    return sections


def _parse_txt(path: Path) -> list[ParsedSection]:
    text = path.read_text(encoding="utf-8", errors="replace").strip()
    if not text:
        return []
    return [ParsedSection(text=text)]
